from docx import Document
import mammoth
from typing import Any
from .base import FileProcessor


class WordProcessor(FileProcessor):
    """Word 文档处理器"""

    def extract_text(self, file_path: str) -> str:
        """提取 Word 文本"""
        doc = Document(file_path)
        return "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])

    def extract_markdown(self, file_path: str) -> str:
        """Word → Markdown（使用 mammoth）"""
        with open(file_path, "rb") as f:
            result = mammoth.convert_to_markdown(f)
            return result.value

    def extract_structured(self, file_path: str) -> dict[str, Any]:
        """提取 Word 结构化数据"""
        doc = Document(file_path)
        return {
            "paragraphs": [p.text for p in doc.paragraphs if p.text.strip()],
            "tables": [[[cell.text for cell in row.cells] for row in table.rows] for table in doc.tables],
        }

    def get_metadata(self, file_path: str) -> dict[str, Any]:
        """获取 Word 元数据"""
        doc = Document(file_path)
        return {
            "page_count": len(doc.sections),
            "paragraph_count": len(doc.paragraphs),
        }
