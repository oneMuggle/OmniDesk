"""chat 接口附件上传测试。

验证契约：
1. multipart 上传 docx → 附件文本注入到 conversation_history 的 system 消息
2. 上传不支持格式（.doc 旧版 OLE）→ 400 + 中文不支持提示
3. 无附件时仍走 JSON 请求
4. 超过 10MB 的文件早期拒绝(400 + 10MB 提示)
5. 无 text/markdown 抽取内容时拒绝(400 + 提取提示)
6. tool_context 携带附件元数据(filename 等)
"""
import io
from unittest.mock import patch

import pytest
from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from rest_framework.test import APIClient


def _make_docx() -> bytes:
    from docx import Document

    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("附件里的合同条款")
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def user(db):
    User = get_user_model()
    return User.objects.create_user(username="tester", password="x")


@pytest.fixture
def client(user):
    c = APIClient()
    c.force_authenticate(user=user)
    return c


class TestChatAttachment:
    @patch("smart_assistant.agent.orchestrator.AgentOrchestrator.process")
    def test_upload_attachment_injects_into_history(self, mock_process, client):
        mock_process.return_value = {
            "answer": "ok", "intent": "general_chat", "tool_used": None,
            "tool_result": None, "sources": None, "usage": {},
        }
        docx = SimpleUploadedFile(
            "合同.docx", _make_docx(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        resp = client.post(
            "/api/smart-assistant/chat/",
            {"query": "合同里写了什么？", "attachment": docx},
            format="multipart",
        )
        assert resp.status_code == 200
        args, kwargs = mock_process.call_args
        # conversation_history 是位置参数(args[1]),tool_context 仅以 kwarg 传入
        history = args[1]
        assert any(
            m.get("role") == "system" and "附件里的合同条款" in m.get("content", "")
            for m in history
        )
        assert "tool_context" in kwargs

    def test_invalid_extension_rejected(self, client):
        bad = SimpleUploadedFile("旧版.doc", b"\xd0\xcf\x11\xe0", content_type="application/msword")
        resp = client.post(
            "/api/smart-assistant/chat/",
            {"query": "x", "attachment": bad},
            format="multipart",
        )
        assert resp.status_code == 400
        assert "不支持" in resp.data["detail"]

    def test_no_attachment_still_json(self, client):
        with patch("smart_assistant.agent.orchestrator.AgentOrchestrator.process") as mock_process:
            mock_process.return_value = {
                "answer": "ok", "intent": "general_chat", "tool_used": None,
                "tool_result": None, "sources": None, "usage": {},
            }
            resp = client.post("/api/smart-assistant/chat/", {"query": "你好"}, format="json")
        assert resp.status_code == 200

    def test_10mb_size_limit_rejected_early(self, client):
        """11MB 文件在 view 层早期拒绝(400 + 10MB 提示),不进入 OfficeExtractor。"""
        big = SimpleUploadedFile(
            "big.docx",
            b"x" * (11 * 1024 * 1024),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        with patch(
            "smart_assistant.extractors.office_extractor.OfficeExtractor.extract"
        ) as mock_extract:
            resp = client.post(
                "/api/smart-assistant/chat/",
                {"query": "x", "attachment": big},
                format="multipart",
            )
        assert resp.status_code == 400
        assert "10MB" in resp.data["detail"]
        # 早期拒绝:OfficeExtractor.extract 不应被调用
        mock_extract.assert_not_called()

    @patch("smart_assistant.extractors.office_extractor.OfficeExtractor.extract")
    def test_empty_text_rejected(self, mock_extract, client):
        """无 text/markdown 抽取内容时返回 400 + 提取提示。"""
        from smart_assistant.extractors.office_extractor import ExtractedDocument

        mock_extract.return_value = ExtractedDocument(
            text="", markdown="", tables=[], sheets=[], metadata={}, format="txt",
        )
        # 用非空字节(避免序列化器直接拒空文件),但 mock 抽取结果为空文本
        blank = SimpleUploadedFile("空白.txt", b"   ", content_type="text/plain")
        resp = client.post(
            "/api/smart-assistant/chat/",
            {"query": "x", "attachment": blank},
            format="multipart",
        )
        assert resp.status_code == 400
        assert "提取到文本" in resp.data["detail"]

    @patch("smart_assistant.agent.orchestrator.AgentOrchestrator.process")
    def test_tool_context_attachment_passed(self, mock_process, client):
        """tool_context.attachment 携带 filename 供下游工具按需取用。"""
        mock_process.return_value = {
            "answer": "ok", "intent": "general_chat", "tool_used": None,
            "tool_result": None, "sources": None, "usage": {},
        }
        docx = SimpleUploadedFile(
            "合同.docx", _make_docx(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        resp = client.post(
            "/api/smart-assistant/chat/",
            {"query": "总结", "attachment": docx},
            format="multipart",
        )
        assert resp.status_code == 200
        _, kwargs = mock_process.call_args
        tool_context = kwargs["tool_context"]
        assert tool_context.attachment is not None
        assert tool_context.attachment["filename"] == "合同.docx"
