"""smart_assistant/tools/office_generate_tool.py — 生成 .docx 文档（require_confirmation）

dry_run 阶段：LLM 规划文档结构 + 变量 → 返回 draft。
confirmed 阶段：用 python-docx 按 structure 构建标题/段落/表格 + 变量替换，
写到临时目录，返回 file_download 卡片信息。
"""

from __future__ import annotations

import json
import re

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from llm_service.router import get_router

from ..cache import get_confirmation_draft
from ..tools_io import create_download_token, save_tmp_office_file
from .base import BaseTool

from observability import get_logger

logger = get_logger(__name__, "smart_assistant")

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

STRUCTURE_PROMPT = """你是文档结构规划器。根据用户描述，规划一份 .docx 文档的结构。

用户描述：{query}

请返回严格 JSON，格式：
{{
  "structure": [
    {{"type": "heading", "content": "标题文本", "level": 1}},
    {{"type": "paragraph", "content": "段落文本，可用 {{变量名}} 占位"}},
    {{"type": "table", "headers": ["列1", "列2"], "rows": [["a", "b"]]}}
  ],
  "variables": {{"变量名": "示例值"}}
}}

要求：
1. structure 每项 type 仅限 heading / paragraph / table
2. 变量用 {{name}} 占位，variables 给出示例值
3. 只输出 JSON，不要其他文字"""


def _plan_document_structure(query: str) -> dict | None:
    """调 LLM 规划文档结构。失败/非法 JSON 返回 None。"""
    try:
        answer, _usage = get_router().generate(prompt=STRUCTURE_PROMPT.format(query=query))
    except Exception as exc:
        logger.warning("文档结构规划 LLM 失败: %s", exc)
        return None
    match = re.search(r"\{.*\}", answer or "", re.DOTALL)
    if not match:
        return None
    try:
        data = json.loads(match.group(0))
    except json.JSONDecodeError:
        return None
    if not isinstance(data.get("structure"), list) or not isinstance(data.get("variables"), dict):
        return None
    return data  # type: ignore[no-any-return]


def _fill(content: str, variables: dict) -> str:
    """把 {name} 占位符替换为变量值。"""
    for key, value in variables.items():
        content = content.replace("{" + key + "}", str(value))
    return content


def _docx_bytes(doc: object) -> bytes:
    import io

    buf = io.BytesIO()
    doc.save(buf)  # type: ignore[attr-defined]
    return buf.getvalue()


def _render_docx_to_file(structure: list, variables: dict, title: str) -> tuple[str, bytes]:
    """用 python-docx 按 structure 构建文档，返回 (相对路径, bytes)。"""
    doc = DocxDocument()
    for item in structure:
        typ = item.get("type")
        content = item.get("content", "")
        if typ == "heading":
            level = int(item.get("level", 1))
            heading = doc.add_heading(level=level)
            run = heading.add_run(_fill(content, variables))
            run.font.size = Pt(16 if level == 1 else 13)
        elif typ == "paragraph":
            doc.add_paragraph(_fill(content, variables))
        elif typ == "table":
            headers = [str(h) for h in item.get("headers", [])]
            rows = [[str(c) for c in r] for r in item.get("rows", [])]
            if headers:
                table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                for j, h in enumerate(headers):
                    table.cell(0, j).text = h
                for i, row in enumerate(rows, start=1):
                    for j, cell in enumerate(row[: len(headers)]):
                        table.cell(i, j).text = cell
    if doc.paragraphs:
        doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    content_bytes = _docx_bytes(doc)
    relative_path = save_tmp_office_file(f"{title}.docx", content_bytes)
    return relative_path, content_bytes


class OfficeGenerateTool(BaseTool):
    """根据用户描述生成 .docx 文档（需二次确认）。"""

    name = "office_generate"
    description = "根据用户描述的结构和变量生成 .docx 文档下载"
    intent_type = "office_generate"
    risk_level = "write"
    require_confirmation = True

    @classmethod
    def get_openai_tool_schema(cls) -> dict:
        """OpenAI strict mode tool schema — 生成 .docx 文档(写操作,需确认)。

        risk_level=write + require_confirmation=True → description 必须显式
        提示需要用户确认,避免 LLM 把写操作当成只读查询使用。
        """
        return {
            "type": "function",
            "function": {
                "name": cls.intent_type,
                "description": (
                    "根据用户描述生成 .docx 文档(写操作,需要用户确认)。"
                    "执行后会先 dry_run 返回文档结构 draft,用户确认后才真正生成文件。"
                    "示例 query: '生成一份设备验收报告模板'、'写一份会议纪要'。"
                ),
                "parameters": {
                    "type": "object",
                    "properties": {
                        "query": {
                            "type": "string",
                            "description": "文档描述,LLM 据此规划 heading/paragraph/table 结构",
                        },
                        "structure_hint": {
                            "type": "string",
                            "description": "可选,用户指定结构关键词(如 '含表格'、'三级标题')",
                        },
                    },
                    "required": ["query"],
                    "additionalProperties": False,
                },
                "strict": True,
            },
        }

    def execute(self, query=None, context=None, **kwargs) -> dict:
        ctx = context if isinstance(context, dict) else {}
        user_id = self._extract_user_id(context)
        if ctx.get("dry_run"):
            return self._dry_run(query, ctx)
        if ctx.get("confirmed"):
            return self._confirmed(query, ctx, user_id)
        return {"found": False, "message": "工具执行异常：未进入 dry_run 或 confirmed 模式"}

    @staticmethod
    def _extract_user_id(context) -> int | str | None:
        """从 ToolContext 或 dict 中提取当前用户 pk，用于把下载 token 绑定到签发者。"""
        if context is None:
            return None
        # ToolContext 对象有 .user 属性；dict 风格 ctx 有 'user' 键
        user = getattr(context, "user", None)
        if user is None and isinstance(context, dict):
            user = context.get("user")
        if user is None:
            return None
        pk = getattr(user, "pk", None)
        if pk is None and isinstance(user, dict):
            pk = user.get("pk") or user.get("id")
        return pk

    def _dry_run(self, query, ctx) -> dict:
        planned = _plan_document_structure(query or "")
        if not planned:
            return {"found": False, "message": "无法根据描述规划文档结构，请补充更明确的要求"}
        structure = planned["structure"]
        title = self._extract_title(structure)
        summary = f"确认生成文档《{title}.docx》"
        return {
            "found": True,
            "draft": {
                "summary": summary,
                "fields": {"structure": structure, "variables": planned["variables"]},
            },
        }

    @staticmethod
    def _extract_title(structure: list, query: str = "") -> str:
        """从结构中提取标题：取第一个 heading 的 content，否则从 query 推断，最后 fallback "文档"。
        忽略 paragraph/table 的 content（它们是正文，不是标题）。
        """
        for item in structure:
            if item.get("type") == "heading":
                content = str(item.get("content", "")).strip()
                if content:
                    return content
        # 从 query 中提取文档类型关键词（去掉"生成""创建"等动词）
        if query:
            stripped = re.sub(
                r"^(生成|创建|写|制作|出一份|来一份|帮我|请|要)",
                "",
                query.strip(),
            ).strip()
            if stripped:
                return stripped
        return "文档"

    def _confirmed(self, query, ctx, user_id=None) -> dict:
        draft = ctx.get("draft")
        if not draft:
            token = ctx.get("confirm_token")
            if token:
                entry = get_confirmation_draft(token)
                if entry and isinstance(entry, dict):
                    draft = entry.get("draft", {}).get("fields")
        if not draft or not draft.get("structure"):
            return {"found": False, "message": "缺少确认时的文档结构，请重新发起生成"}
        structure = draft["structure"]
        variables = draft.get("variables", {})
        title = self._extract_title(structure, query or "")
        if user_id is None:
            logger.error("OfficeGenerateTool 缺少签发用户 ID,拒绝签发下载 token")
            return {"found": False, "message": "无法签发下载链接，请重新登录后再试"}
        try:
            relative_path, _bytes = _render_docx_to_file(structure, variables, title)
            token = create_download_token(relative_path, user_id)
        except Exception as exc:
            logger.exception("生成 docx 失败")
            return {"found": False, "message": "生成文档失败，请稍后重试"}
        return {
            "found": True,
            "summary": f"文档《{title}.docx》已生成",
            "file_download": {
                "filename": f"{title}.docx",
                "download_url": f"/api/smart-assistant/office-download/{token}/",
                "content_type": DOCX_MIME,
            },
        }
