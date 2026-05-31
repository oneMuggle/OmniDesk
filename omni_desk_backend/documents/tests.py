from django.test import TestCase
from django.contrib.auth import get_user_model
from rest_framework.test import APITestCase, APIClient
from rest_framework import status
from django.urls import reverse
from .models import DocumentTemplate, GeneratedDocument, Book, Chapter, Comment, Annotation, Tag, EBook
from projects.models import Project

class DocumentModelsTest(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(username='doc_user', password='password')
        self.project = Project.objects.create(name='Test Project', description='A project for testing.')

    def test_document_template_str(self):
        template = DocumentTemplate.objects.create(
            project=self.project,
            name='Test Template',
            template_type='tech_design',
            content='Template content',
            owner=self.user
        )
        self.assertEqual(str(template), '技术方案文档 - Test Template')

    def test_generated_document_str(self):
        template = DocumentTemplate.objects.create(name='GenDoc Template', owner=self.user, template_type='test_case')
        gen_doc = GeneratedDocument.objects.create(
            template=template,
            content='Generated content',
            generated_by=self.user
        )
        self.assertIn(template.name, str(gen_doc))

    def test_tag_str(self):
        tag = Tag.objects.create(name='Python')
        self.assertEqual(str(tag), 'Python')

    def test_book_str(self):
        book = Book.objects.create(title='My First Book')
        self.assertEqual(str(book), 'My First Book')

    def test_chapter_str(self):
        book = Book.objects.create(title='Book for Chapter')
        chapter = Chapter.objects.create(book=book, title='Chapter 1', order=1, content_md='md')
        self.assertEqual(str(chapter), f'{book.title} - 1. Chapter 1')

    def test_comment_str(self):
        book = Book.objects.create(title='Book for Comment')
        chapter = Chapter.objects.create(book=book, title='Chapter for Comment', order=1, content_md='md')
        comment = Comment.objects.create(chapter=chapter, user=self.user, content='A comment')
        self.assertEqual(str(comment), f'Comment by {self.user} on {chapter}')

    def test_annotation_str(self):
        book = Book.objects.create(title='Book for Annotation')
        chapter = Chapter.objects.create(book=book, title='Chapter for Annotation', order=1, content_md='md')
        annotation = Annotation.objects.create(chapter=chapter, user=self.user, selected_text='highlight')
        self.assertEqual(str(annotation), f'Annotation by {self.user} on {chapter}')

    def test_ebook_str(self):
        ebook = EBook.objects.create(title='My EBook', content='Ebook content')
        self.assertEqual(str(ebook), 'My EBook')

class DocumentViewSetTests(APITestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(username='api_doc_user', password='password', is_staff=True, is_superuser=True)
        self.client = APIClient()
        self.client.force_authenticate(user=self.user)
        self.project = Project.objects.create(name='API Test Project', description='...')

    def test_create_document_template(self):
        url = reverse('document-template-list')
        data = {
            'name': 'API Template',
            'template_type': 'meeting_minutes',
            'content': 'Meeting content',
            'project': self.project.id
        }
        response = self.client.post(url, data, format='multipart')
        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        self.assertEqual(DocumentTemplate.objects.count(), 1)
        self.assertEqual(DocumentTemplate.objects.first().name, 'API Template')

    def test_list_books_filtered_by_project(self):
        Book.objects.all().delete() # Clean up before test
        book1 = Book.objects.create(title='Project Book', project=self.project)
        book2 = Book.objects.create(title='Other Book')
        url = reverse('book-list')
        response = self.client.get(url, {'project_id': self.project.id})
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(len(response.data['results']), 1)
        self.assertEqual(response.data['results'][0]['title'], 'Project Book')

    def test_add_comment_to_chapter(self):
        book = Book.objects.create(title='Commentable Book')
        chapter = Chapter.objects.create(book=book, title='Chapter', order=1, content_md='md')
        url = reverse('chapter-add-comment', args=[chapter.id])
        data = {'content': 'This is a test comment.'}
        response = self.client.post(url, data, format='json')
        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        self.assertEqual(chapter.comments.count(), 1)

import zipfile
import tempfile
from django.core.files.uploadedfile import SimpleUploadedFile
from django.conf import settings
import os
import shutil
import importlib.util
import unittest

@unittest.skipIf(
    importlib.util.find_spec('taggit') is None,
    'django-taggit not installed (optional dependency)',
)
class BookImportViewTests(APITestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(username='import_user', password='password', is_staff=True, is_superuser=True)
        self.client = APIClient()
        self.client.force_authenticate(user=self.user)
        self.url = reverse('import-book')

    def tearDown(self):
        # Clean up media files created during tests
        book_images_path = os.path.join(settings.MEDIA_ROOT, 'book_images')
        if os.path.exists(book_images_path):
            shutil.rmtree(book_images_path)
        covers_path = os.path.join(settings.MEDIA_ROOT, 'covers')
        if os.path.exists(covers_path):
            shutil.rmtree(covers_path)

    def test_import_book_from_zip(self):
        """Test importing a book from a zip file containing markdown and an image."""
        # Create a temporary directory to build the zip file
        with tempfile.TemporaryDirectory() as temp_dir:
            # Create a dummy markdown file
            md_content = "# Test Zip Book\n\nThis is the content with an image ![alt text](test_image.png)."
            md_file_path = os.path.join(temp_dir, 'book.md')
            with open(md_file_path, 'w', encoding='utf-8') as f:
                f.write(md_content)

            # Create a dummy image file
            image_content = b'GIF89a\x01\x00\x01\x00\x80\x00\x00\xff\xff\xff\x00\x00\x00!\xf9\x04\x01\x00\x00\x00\x00,\x00\x00\x00\x00\x01\x00\x01\x00\x00\x02\x02D\x01\x00;'
            image_file_path = os.path.join(temp_dir, 'test_image.png')
            with open(image_file_path, 'wb') as f:
                f.write(image_content)

            # Create the zip file
            zip_file_path = os.path.join(temp_dir, 'book.zip')
            with zipfile.ZipFile(zip_file_path, 'w') as zf:
                zf.write(md_file_path, 'book.md')
                zf.write(image_file_path, 'test_image.png')

            # Prepare the upload
            with open(zip_file_path, 'rb') as f:
                zip_file = SimpleUploadedFile('book.zip', f.read(), content_type='application/zip')

            data = {'file': zip_file, 'title': 'My Zip Imported Book'}
            response = self.client.post(self.url, data, format='multipart')

            self.assertEqual(response.status_code, status.HTTP_201_CREATED)
            self.assertTrue(Book.objects.filter(title='My Zip Imported Book').exists())
            book = Book.objects.get(title='My Zip Imported Book')
            self.assertEqual(book.chapters.count(), 1)
            chapter = book.chapters.first()
            self.assertIn('Test Zip Book', chapter.title)
            # Check if image path was correctly updated in the content
            self.assertIn(f'{settings.MEDIA_URL}book_images/My_Zip_Imported_Book/test_image.png', chapter.content_md)


from unittest.mock import patch, MagicMock
from pathlib import Path
import pypdf
from docx import Document as DocxDocument
from .file_processing import (
    process_uploaded_file,
    extract_text_from_pdf,
    extract_text_from_docx,
    call_mineru_ocr,
)


class FileProcessingTests(TestCase):
    def setUp(self):
        # Create a temporary directory
        self.temp_dir_obj = tempfile.TemporaryDirectory()
        self.temp_dir = self.temp_dir_obj.name

    def tearDown(self):
        # Cleanup the directory
        self.temp_dir_obj.cleanup()

    @patch('documents.file_processing.extract_text_from_docx')
    def test_process_uploaded_file_docx_success(self, mock_extract):
        """Test processing a DOCX file successfully."""
        mock_extract.return_value = "Hello from DOCX"
        file_content = b"dummy docx"
        uploaded_file = SimpleUploadedFile("test.docx", file_content)
        
        text = process_uploaded_file(uploaded_file, self.temp_dir)
        
        self.assertTrue(os.path.exists(os.path.join(self.temp_dir, "test.docx")))
        mock_extract.assert_called_once()
        self.assertEqual(text, "Hello from DOCX")

    @patch('documents.file_processing.extract_text_from_pdf')
    @patch('documents.file_processing.call_mineru_ocr')
    def test_process_uploaded_file_pdf_fallback(self, mock_ocr, mock_extract_pdf):
        """Test PDF processing falls back to OCR on failure."""
        mock_extract_pdf.return_value = None
        mock_ocr.return_value = "Hello from OCR"
        file_content = b"dummy pdf"
        uploaded_file = SimpleUploadedFile("test.pdf", file_content)

        text = process_uploaded_file(uploaded_file, self.temp_dir)

        mock_extract_pdf.assert_called_once()
        mock_ocr.assert_called_once()
        self.assertEqual(text, "Hello from OCR")

    @patch('documents.file_processing.call_mineru_ocr')
    def test_process_uploaded_file_image(self, mock_ocr):
        """Test image file processing calls OCR."""
        mock_ocr.return_value = "Hello from Image OCR"
        file_content = b"dummy image"
        uploaded_file = SimpleUploadedFile("test.png", file_content)

        text = process_uploaded_file(uploaded_file, self.temp_dir)
        mock_ocr.assert_called_once()
        self.assertEqual(text, "Hello from Image OCR")

    def test_process_uploaded_file_unknown_type_text(self):
        """Test processing an unknown file type as text."""
        file_content = "Just a text file".encode('utf-8')
        uploaded_file = SimpleUploadedFile("test.txt", file_content)

        text = process_uploaded_file(uploaded_file, self.temp_dir)
        self.assertEqual(text, "Just a text file")

    def test_process_uploaded_file_unknown_type_binary(self):
        """Test processing a binary file that cannot be decoded."""
        file_content = b'\x80\x81\x82' # Invalid utf-8
        uploaded_file = SimpleUploadedFile("test.bin", file_content)

        text = process_uploaded_file(uploaded_file, self.temp_dir)
        self.assertEqual(text, "")

    def test_extract_text_from_pdf_success(self):
        """Test successful text extraction from PDF."""
        with patch('pypdf.PdfReader') as mock_reader:
            mock_page = MagicMock()
            mock_page.extract_text.return_value = "PDF page text. "
            mock_instance = mock_reader.return_value
            mock_instance.pages = [mock_page, mock_page]
            
            dummy_path = os.path.join(self.temp_dir, 'dummy.pdf')
            Path(dummy_path).touch()
            
            text = extract_text_from_pdf(dummy_path)
            self.assertEqual(text, "PDF page text. PDF page text. ")

    def test_extract_text_from_pdf_error(self):
        """Test text extraction failure from PDF."""
        with patch('pypdf.PdfReader', side_effect=pypdf.errors.PdfReadError):
            dummy_path = os.path.join(self.temp_dir, 'dummy.pdf')
            Path(dummy_path).touch()
            text = extract_text_from_pdf(dummy_path)
            self.assertIsNone(text)

    def test_extract_text_from_docx_success(self):
        """Test successful text extraction from DOCX."""
        doc = DocxDocument()
        doc.add_paragraph("DOCX paragraph.")
        doc.add_paragraph("Another paragraph.")
        dummy_path = os.path.join(self.temp_dir, 'dummy.docx')
        doc.save(dummy_path)

        text = extract_text_from_docx(dummy_path)
        self.assertEqual(text, "DOCX paragraph.\nAnother paragraph.\n")

    @patch('requests.post')
    def test_call_mineru_ocr_success(self, mock_post):
        """Test successful call to Mineru OCR."""
        mock_response = MagicMock()
        mock_response.raise_for_status.return_value = None
        mock_response.json.return_value = {'text': 'ocr text'}
        mock_post.return_value = mock_response

        dummy_path = os.path.join(self.temp_dir, 'dummy.png')
        Path(dummy_path).touch()

        with self.settings(MINERU_API_URL='http://fake.url', MINERU_API_KEY='fake-key'):
            text = call_mineru_ocr(dummy_path)
        
        self.assertEqual(text, 'ocr text')
        mock_post.assert_called_once()

    def test_call_mineru_ocr_no_settings(self):
        """Test Mineru OCR call raises error if settings are missing."""
        dummy_path = os.path.join(self.temp_dir, 'dummy.png')
        Path(dummy_path).touch()

        with self.settings(MINERU_API_URL=None):
            with self.assertRaises(ValueError):
                call_mineru_ocr(dummy_path)
