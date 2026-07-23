from docx import Document
import os

os.makedirs('tests/fixtures', exist_ok=True)

doc = Document()
doc.add_heading('测试文档', 1)
doc.add_paragraph('这是第一段内容。')
doc.add_paragraph('这是第二段内容。')

table = doc.add_table(rows=2, cols=2)
table.cell(0, 0).text = '姓名'
table.cell(0, 1).text = '年龄'
table.cell(1, 0).text = '张三'
table.cell(1, 1).text = '25'

doc.save('tests/fixtures/sample.docx')
print("Word test file created")
